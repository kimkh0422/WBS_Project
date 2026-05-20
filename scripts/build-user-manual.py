"""
지엠티 프로젝트 매니저 사용자 매뉴얼 (Word 파일) 생성 스크립트.

실행:
  python scripts/build-user-manual.py
출력:
  docs/manual/사용자매뉴얼.docx

이미지 자리에는 회색 박스 + 안내 텍스트가 들어가며, 사용자가 직접 캡처해
붙여넣을 수 있게 한다. 각 이미지 자리는 번호와 "무엇을 캡처해야 하는지"
설명을 포함한다.
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ────────────────────────────────────────────────────────────────────────────
# 헬퍼
# ────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str) -> None:
    """표 셀 배경색 설정 (XML 직접 조작)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_image_placeholder(doc: Document, image_no: int, what_to_capture: str) -> None:
    """이미지 붙여넣기 자리.

    회색 박스(1×1 표)에 [📷 이미지 N] 라벨 + 캡처 안내 텍스트.
    사용자는 박스 안에 이미지를 붙여넣은 뒤 안내 텍스트를 지운다.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, "F2F2F2")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 라벨
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"📷 이미지 {image_no}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 안내
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(what_to_capture)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # 표 너비를 본문 너비에 맞춤 (대략)
    for row in table.rows:
        for c in row.cells:
            tc_pr = c._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), "9000")
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)

    # 표 아래에 빈 줄
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
        # 한글 폰트 명시
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "맑은 고딕"
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        # 들여쓰기로 단계 표시 (스타일 없이)
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    r = p.runs[0] if p.runs else p.add_run("")
    p.runs[0].text = text if not p.runs[0].text else p.runs[0].text  # noop
    # 새 run으로 텍스트 채우기
    if not p.runs[0].text:
        p.runs[0].text = text
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_step(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r1 = p.add_run(f"{n}. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    for run in (r1, r2):
        run.font.name = "맑은 고딕"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_tip(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, "FFF8E1")
    p = cell.paragraphs[0]
    r = p.add_run(f"💡 {text}")
    r.font.size = Pt(9)
    r.font.name = "맑은 고딕"
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()


def add_kbd_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "단축키"
    hdr[1].text = "동작"
    for c in hdr:
        set_cell_bg(c, "E8F0FE")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = "맑은 고딕"
    for i, (k, v) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = "맑은 고딕"
                    rPr = r._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.append(rFonts)
                    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()


# ────────────────────────────────────────────────────────────────────────────
# 문서 작성
# ────────────────────────────────────────────────────────────────────────────

def build() -> Path:
    doc = Document()

    # 본문 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── 표지 ──────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("\n\n\n지엠티 프로젝트 매니저\n사용자 매뉴얼")
    rt.bold = True
    rt.font.size = Pt(28)
    rt.font.name = "맑은 고딕"
    rPr = rt._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("\n다수 프로젝트의 WBS를 한 화면에서 관리·시각화·공유\n\n")
    rs.font.size = Pt(12)
    rs.italic = True
    rs.font.name = "맑은 고딕"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("v1.0 · 2026-05-07")
    rm.font.size = Pt(10)
    rm.font.name = "맑은 고딕"

    doc.add_page_break()

    # ── 목차 안내 ─────────────────────────────────────────────────────────
    add_heading(doc, "목차", level=1)
    toc_lines = [
        "1. 프로그램 소개",
        "2. 시작하기 (회원가입·로그인)",
        "3. 프로젝트 만들기",
        "4. 작업(WBS) 추가하기",
        "5. 표에서 셀 편집하기",
        "6. 트리 구조 만들기 (들여쓰기/내어쓰기)",
        "7. 다중 선택과 일괄 수정",
        "8. 6가지 화면(시점) 전환",
        "9. 필터·정렬·검색",
        "10. 프로젝트 공유와 권한 부여",
        "11. 환경설정",
        "12. 데이터 가져오기·내보내기·백업",
        "13. AI 보조와 주간보고",
        "14. 키보드 단축키 일람",
        "15. 자주 묻는 질문",
    ]
    for line in toc_lines:
        add_para(doc, line, size=10)
    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 1. 프로그램 소개
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "1. 프로그램 소개", level=1)
    add_para(
        doc,
        "지엠티 프로젝트 매니저는 다수 프로젝트의 작업 분해 구조(WBS)를 단일 화면에서 "
        "관리·시각화·공유하는 협업형 웹 애플리케이션입니다. 같은 데이터를 6가지 시점"
        "(표·표+간트·간트만·칸반·마인드맵·대시보드)으로 볼 수 있어, 사용자는 화면만 "
        "전환하며 다른 관점에서 일정을 점검할 수 있습니다.",
    )
    add_heading(doc, "1.1 핵심 가치", level=2)
    add_bullet(doc, "여러 프로젝트의 진행 상황을 한 화면에서 통합 조회")
    add_bullet(doc, "표·간트·칸반·마인드맵 시점 전환으로 다양한 사고 방식 지원")
    add_bullet(doc, "시작일·공수만 입력하면 휴일·의존성을 반영해 종료일 자동 산정")
    add_bullet(doc, "1초 디바운스 자동 저장 + Undo 50회 + JSON 백업으로 실수 회복")
    add_bullet(doc, "조직(부서·인원) 데이터 연동으로 담당자 자동완성")

    add_image_placeholder(
        doc, 1,
        "전체 화면 — 헤더(좌측 로고·프로젝트 드롭다운, 가운데 뷰 탭, 우측 [+ 새 작업]·필터·메뉴)와 "
        "하단 작업 목록(표 또는 표+간트)이 모두 보이도록 캡처."
    )

    add_heading(doc, "1.2 권한 체계 한눈에 보기", level=2)
    add_para(doc, "프로젝트마다 한 명의 소유자(만든 사람)와 시스템 관리자만 편집할 수 있습니다. 그 외 회원은 보기만 가능합니다.")
    add_bullet(doc, "비승인 회원: 가입했지만 관리자가 아직 승인하지 않은 상태. 로컬에서만 시범 사용 가능.")
    add_bullet(doc, "승인 회원: 모든 프로젝트를 조회할 수 있으나 편집은 본인 소유 또는 권한 부여 받은 프로젝트에서만.")
    add_bullet(doc, "프로젝트 소유자: 그 프로젝트에 한해 모든 편집·공유·삭제 권한 보유.")
    add_bullet(doc, "시스템 관리자: 모든 프로젝트·회원·환경설정 변경 가능.")
    add_bullet(doc, "조직 책임자: 본인 산하 부서 회원의 승인·권한 변경 가능.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 2. 시작하기
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "2. 시작하기 (회원가입·로그인)", level=1)
    add_heading(doc, "2.1 회원가입", level=2)
    add_step(doc, 1, "초기 화면에서 [회원가입] 클릭.")
    add_step(doc, 2, "회사 이메일·비밀번호를 입력하고 [가입].")
    add_step(doc, 3, "메일함에서 인증 메일을 확인합니다(필요 시).")
    add_step(doc, 4, "관리자가 승인할 때까지 대기 — 승인 전에는 클라우드 동기화가 막혀 있고, 로컬 시범 사용만 가능합니다.")

    add_image_placeholder(
        doc, 2,
        "로그인/회원가입 화면 — 이메일·비밀번호 입력란과 [로그인]·[회원가입] 버튼이 보이도록 캡처."
    )

    add_tip(doc, "관리자에게 회원가입 사실을 알리면 알림 벨로 즉시 확인되어 빠르게 승인됩니다.")

    add_heading(doc, "2.2 로그인 후 첫 화면", level=2)
    add_para(doc, "로그인하면 프로젝트가 0개일 경우 빈 화면 안내가 표시되고, 기존 프로젝트가 있다면 마지막에 보던 프로젝트가 자동 선택됩니다.")
    add_bullet(doc, "헤더 좌측: 프로젝트 드롭다운 (즐겨찾기 + 그룹별 + 검색)")
    add_bullet(doc, "헤더 가운데: 6가지 뷰 탭(대시보드 / 투입현황 / 표+간트 / 표만 / 간트만 / 칸반 / 마인드맵 / 프로젝트)")
    add_bullet(doc, "헤더 우측: 필터 토글 / Undo·Redo / 더보기(⋮) / [+ 새 작업] / 사용자 메뉴")

    add_image_placeholder(
        doc, 3,
        "헤더 영역 확대 — 프로젝트 드롭다운, 뷰 탭, [+ 새 작업] 버튼, 사용자 메뉴가 잘 보이도록 캡처."
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 3. 프로젝트 만들기
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "3. 프로젝트 만들기", level=1)
    add_heading(doc, "3.1 새 프로젝트 생성", level=2)
    add_step(doc, 1, "헤더 [+ 프로젝트] 또는 더보기(⋮) → [프로젝트 추가] 클릭.")
    add_step(doc, 2, "필수 입력: 프로젝트 이름.")
    add_step(doc, 3, "선택 입력: 설명 / 시작일·종료일 / 프로젝트 PM.")
    add_step(doc, 4, "주간보고용 메타(과제명 약어·전체과제명·구분·주관기관·예산·전체기간) 입력.")
    add_step(doc, 5, "투입인원 추가: [+ 인원 추가] 후 이름 입력. 입력란에서 Enter를 누르면 다음 인원이 자동으로 추가됩니다.")
    add_step(doc, 6, "[저장] 클릭.")

    add_image_placeholder(
        doc, 4,
        "프로젝트 추가/수정 모달 — 필수 입력(프로젝트 이름) + 기본 정보(시작일·종료일) + 투입인원 영역이 보이도록 캡처."
    )

    add_tip(doc, "투입인원의 담당자 입력란을 클릭하면 조직 회원 전체 목록이 자동완성으로 표시됩니다. '김'을 입력하면 김씨가 들어간 이름만 좁혀집니다.")

    add_image_placeholder(
        doc, 5,
        "투입인원 입력란에 '김'을 친 직후 — 자동완성 드롭다운에 '김XX (부서 · 직위)' 형식의 후보가 보이는 화면."
    )

    add_heading(doc, "3.2 프로젝트 일정 변경", level=2)
    add_para(doc, "프로젝트의 시작일·종료일을 줄이면 그 범위 밖의 작업은 자동으로 클램프(범위 안으로 잘림)됩니다. 늘리는 것은 작업에 영향이 없습니다.")
    add_bullet(doc, "변경 후 프로젝트 모달을 다시 열면 새 값이 그대로 보입니다.")
    add_bullet(doc, "표 상단 SummaryBar의 '기간' 표시에도 프로젝트의 새 일정이 반영됩니다.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 4. 작업 추가
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "4. 작업(WBS) 추가하기", level=1)
    add_para(doc, "작업은 여러 진입점에서 추가할 수 있습니다. 가장 흔한 4가지를 순서대로 안내합니다.")

    add_heading(doc, "4.1 표 상단 빠른 추가", level=2)
    add_step(doc, 1, "표 상단(컬럼 헤더 바로 아래)에 항상 보이는 청색 입력 행을 클릭.")
    add_step(doc, 2, "작업명을 입력 후 Enter — 새 작업이 루트 레벨로 추가됩니다.")
    add_step(doc, 3, "추가된 행에서 바로 시작일·종료일·공수를 채울 수 있습니다.")

    add_image_placeholder(
        doc, 6,
        "표 상단의 [+ 새 작업 추가 (Enter 키 입력)…] 청색 입력 행 — 컬럼 헤더 바로 아래에 sticky로 고정된 모습."
    )

    add_heading(doc, "4.2 키보드 단축키로 추가", level=2)
    add_kbd_table(doc, [
        ("Enter (작업명 편집 중)", "같은 레벨 형제로 현재 행 아래에 새 작업"),
        ("Shift+Enter", "같은 레벨 형제로 현재 행 위에 새 작업"),
        ("Insert", "기준 행의 하위(자식)에 새 작업"),
        ("Shift+Insert", "기준 행 위에 같은 레벨 새 작업"),
    ])

    add_heading(doc, "4.3 우클릭(컨텍스트 메뉴)", level=2)
    add_step(doc, 1, "표의 작업 행 위에서 우클릭.")
    add_step(doc, 2, "[새 작업 추가 — 위/아래/하위] 중 선택.")

    add_image_placeholder(
        doc, 7,
        "표 행 우클릭 시 나타나는 컨텍스트 메뉴 — '새 작업 추가 위/아래/하위' 항목이 보이도록 캡처."
    )

    add_heading(doc, "4.4 칸반·마인드맵에서 추가", level=2)
    add_bullet(doc, "칸반: 컬럼 하단 [+ 카드 추가] → 제목 입력 → Enter. 해당 컬럼의 상태로 루트 작업 생성.")
    add_bullet(doc, "마인드맵: 빈 상태 [작업 추가] 또는 노드 선택 후 Tab/Ctrl+Enter로 자식 노드 추가.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 5. 셀 편집
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "5. 표에서 셀 편집하기", level=1)
    add_para(doc, "엑셀과 비슷한 2단계 클릭 패턴을 사용합니다.")
    add_step(doc, 1, "셀을 한 번 클릭 — 행이 선택되고 셀에 점선 테두리(포커스).")
    add_step(doc, 2, "같은 셀을 한 번 더 클릭 (또는 F2 / Enter) — 입력 모드 진입.")
    add_step(doc, 3, "값 입력 / 드롭다운 선택 / 날짜 픽커.")
    add_step(doc, 4, "Enter로 확정. 같은 셀에 머무르며 ←/→로 다른 셀로 이동.")
    add_step(doc, 5, "Tab으로 다음 컬럼, Shift+Tab으로 이전 컬럼.")

    add_image_placeholder(
        doc, 8,
        "표에서 시작일 셀을 두 번째 클릭한 후 날짜 입력 모드 — 날짜 picker가 열린 상태."
    )

    add_heading(doc, "5.1 컬럼별 입력 방식", level=2)
    add_bullet(doc, "작업명: 텍스트 자유 입력")
    add_bullet(doc, "시작일·종료일: 날짜 picker 또는 yyyy-mm-dd 직접 입력")
    add_bullet(doc, "공수: 숫자 (소수 1자리). 단위는 프로젝트의 work_effort_unit (일/시간/인월)")
    add_bullet(doc, "진척도: 0~100 (소수 가능)")
    add_bullet(doc, "가중치: 같은 부모 형제 합 100 자동 보장")
    add_bullet(doc, "상태: 환경설정에서 정의된 상태 드롭다운")
    add_bullet(doc, "담당자: 조직 회원 + 프로젝트 등록 인원 자동완성")
    add_bullet(doc, "선행작업: 같은 프로젝트 내 작업 다중 선택 (순환 자동 차단)")

    add_image_placeholder(
        doc, 9,
        "담당자 셀 편집 — 자동완성 드롭다운에 조직 회원 목록이 떠 있는 모습 (이름 + 부서/직위 라벨이 보이게)."
    )

    add_heading(doc, "5.2 자동 일정 산정", level=2)
    add_para(doc, "시작일과 공수만 입력하면 종료일이 휴일·투입율을 반영해 자동으로 계산됩니다. 종료일을 직접 바꾸면 그 값이 잠금되어 자동 변경되지 않습니다.")

    add_tip(doc, "필드 옆 자물쇠 표시는 자동 산정·롤업·AI 보정에서 그 필드를 제외하는 잠금입니다. 사용자가 직접 편집하면 자동으로 잠깁니다.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 6. 트리 만들기
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "6. 트리 구조 만들기 (들여쓰기/내어쓰기)", level=1)
    add_para(doc, "WBS 트리는 부모-자식 관계로 구성됩니다. 작업을 입력한 뒤 Tab으로 들여쓰기하여 부모의 자식으로 만듭니다.")
    add_kbd_table(doc, [
        ("Tab", "현재 행을 직전 형제의 자식으로 들여쓰기"),
        ("Shift+Tab", "현재 행을 부모의 형제로 내어쓰기"),
        ("Alt+↑", "같은 부모 안에서 한 칸 위로 이동"),
        ("Alt+↓", "같은 부모 안에서 한 칸 아래로 이동"),
        ("Shift+→", "트리 펼치기"),
        ("Shift+←", "트리 접기"),
        ("Ctrl+Alt+1~9", "전체 N 레벨까지 펼치기"),
    ])

    add_image_placeholder(
        doc, 10,
        "트리 구조 예시 — P1 아래 W1.1, W1.2, W1.3이 들여쓰기되어 있고, W1.1 아래에는 W1.1.1이 더 들여쓰기된 상태."
    )

    add_tip(doc, "WBS 정렬을 'wbs asc'(기본)이 아닌 다른 컬럼으로 바꾸면 Tab/Shift+Tab/Alt+↑↓가 비활성화됩니다. 트리 표시 순서와 충돌하기 때문입니다.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 7. 다중 선택
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "7. 다중 선택과 일괄 수정", level=1)
    add_heading(doc, "7.1 선택 방법", level=2)
    add_kbd_table(doc, [
        ("체크박스 클릭", "단일 토글"),
        ("Space (포커스 행에서)", "체크박스 토글"),
        ("Shift+클릭", "마지막 선택과 클릭 행 사이 모두 선택"),
        ("Ctrl+클릭", "비연속 추가/제거"),
        ("헤더 체크박스", "현재 표시 중인 작업 모두 선택/해제"),
    ])

    add_heading(doc, "7.2 일괄 수정 바", level=2)
    add_para(doc, "2개 이상 선택하면 화면 하단에 [일괄 수정] 바가 나타납니다.")
    add_bullet(doc, "변경 가능 필드: 상태 / 담당자(자식 캐스케이드 옵션) / 공수 / 진척율 / 가중치 / 시작일 / 종료일")
    add_bullet(doc, "[선행 순차 연결]: 선택한 표시 순서대로 위→아래 행을 선행으로 자동 체인 구성")
    add_bullet(doc, "[선행작업 지우기]: 선택 작업의 모든 의존성 제거 + 자동 재추가 차단")
    add_bullet(doc, "[삭제]: 선택 작업 일괄 삭제 (확인 다이얼로그)")

    add_image_placeholder(
        doc, 11,
        "표에서 작업 3건을 체크한 상태 + 화면 하단의 [일괄 수정] 바가 떠 있는 모습 (담당자·상태·공수 입력란 + [적용]·[선행 순차 연결] 버튼)."
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 8. 6가지 시점
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "8. 6가지 화면(시점) 전환", level=1)
    add_para(doc, "헤더 가운데 탭으로 화면을 전환합니다. 같은 데이터를 다른 시점으로 보여주며, 프로젝트 전환 시에도 현재 시점이 유지됩니다.")

    add_heading(doc, "8.1 표만 (table)", level=2)
    add_para(doc, "엑셀처럼 모든 컬럼을 한눈에. 키보드 중심 빠른 편집에 적합.")
    add_image_placeholder(
        doc, 12,
        "표만 화면 — 컬럼 헤더 + 작업 행들 + 상단 sticky [+ 새 작업 추가] 행이 보이는 캡처."
    )

    add_heading(doc, "8.2 표+간트 (list)", level=2)
    add_para(doc, "좌측 표, 우측 간트 차트. 가운데 핸들로 너비 비율 조정 가능. 표와 간트의 스크롤이 동기화.")
    add_image_placeholder(
        doc, 13,
        "표+간트 화면 — 좌측 표와 우측 간트 막대들이 같은 행으로 정렬된 모습."
    )

    add_heading(doc, "8.3 간트만 (gantt)", level=2)
    add_para(doc, "시간축이 더 넓게 보이는 풀 간트. 막대 드래그로 일정 이동, 좌우 끝 드래그로 길이 조정.")
    add_image_placeholder(
        doc, 14,
        "간트만 화면 — 시간축 헤더 + 작업 막대들 + 의존성 화살표가 보이는 캡처."
    )

    add_heading(doc, "8.4 칸반 (kanban)", level=2)
    add_para(doc, "상태별 컬럼에 카드. 카드 드래그로 상태 변경. 진척도 자동 연동.")
    add_image_placeholder(
        doc, 15,
        "칸반 화면 — 상태별 컬럼(시작 전·진행 중·완료 등)과 각 컬럼의 카드들."
    )

    add_heading(doc, "8.5 마인드맵 (mindmap)", level=2)
    add_para(doc, "방사형 노드 트리. 자유 레이아웃 + 색상으로 구조 시각화.")
    add_image_placeholder(
        doc, 16,
        "마인드맵 화면 — 중앙 노드와 연결된 자식 노드들."
    )

    add_heading(doc, "8.6 대시보드 (dashboard)", level=2)
    add_para(doc, "진행 중·지연·완료 카운트, 이번 주 작업, 프로젝트별 진척도, 번다운 차트 등 종합 지표.")
    add_image_placeholder(
        doc, 17,
        "대시보드 화면 — 카드 위젯들과 이번 주 작업 목록이 보이는 전체 화면."
    )

    add_tip(doc, "대시보드의 카드/버튼을 클릭하면 해당 필터가 자동 적용된 작업 보기로 점프합니다(예: '이번 주 작업' → list 뷰 + 주간 필터).")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 9. 필터·정렬·검색
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "9. 필터·정렬·검색", level=1)
    add_heading(doc, "9.1 필터 바", level=2)
    add_step(doc, 1, "헤더 [필터] 토글 클릭 → 필터 바 표시.")
    add_step(doc, 2, "프로젝트 / 상태 / 담당자 / 기간 / 레벨 / 마일스톤 / 이슈 / 지연 / 이번 주 완료 등으로 좁히기.")
    add_step(doc, 3, "한 번 더 토글하면 필터 OFF (조건은 메모리에 보존, 다음 ON 시 복원).")

    add_image_placeholder(
        doc, 18,
        "필터 바가 펼쳐진 상태 — 프로젝트 / 상태 / 담당자 / 마일스톤 등 chip이 가로로 나열된 모습."
    )

    add_heading(doc, "9.2 정렬", level=2)
    add_bullet(doc, "기본은 WBS 번호 오름차순(트리 순서).")
    add_bullet(doc, "컬럼 헤더를 클릭하면 그 컬럼으로 정렬 (한 번 더 클릭 = 내림차순, 한 번 더 = 해제).")
    add_bullet(doc, "WBS 정렬이 아니면 트리 들여쓰기 동작(Tab 등)이 비활성화됩니다.")

    add_heading(doc, "9.3 검색", level=2)
    add_step(doc, 1, "Ctrl+K 또는 헤더 검색 버튼.")
    add_step(doc, 2, "키워드 입력 — 작업명·메모·담당자·사용자 정의 컬럼에서 검색.")
    add_step(doc, 3, "↑/↓로 이동, Enter로 점프 — 다른 프로젝트 작업도 점프해서 표 화면으로 자동 전환.")

    add_image_placeholder(
        doc, 19,
        "검색 모달 — 검색창에 키워드 입력 + 결과 목록(프로젝트명·WBS 번호·작업명)."
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 10. 프로젝트 공유
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "10. 프로젝트 공유와 권한 부여", level=1)
    add_para(doc, "프로젝트별로 보기/편집 권한을 다른 회원에게 부여할 수 있습니다.")

    add_heading(doc, "10.1 공유 모달 진입", level=2)
    add_step(doc, 1, "헤더 ⋮ → [공유] 또는 프로젝트 카드 ⋮ → [공유].")
    add_step(doc, 2, "현재 멤버 목록 + 가입 대기(미가입 사전 등록) + 사용자 권한 부여 영역이 표시됩니다.")

    add_image_placeholder(
        doc, 20,
        "공유 모달 — 멤버 목록 + 가입 대기 영역 + 사용자 권한 부여 입력 영역이 모두 보이는 캡처."
    )

    add_heading(doc, "10.2 가입 회원 일괄 추가", level=2)
    add_step(doc, 1, "조직 드롭다운에서 부서를 좁힌 뒤 사용자 목록에서 체크.")
    add_step(doc, 2, "역할(보기/편집) 선택 → [선택 추가].")

    add_heading(doc, "10.3 미가입자 사전 등록 (가입 시 자동 권한 부여)", level=2)
    add_para(doc, "아직 회원가입을 하지 않은 사람을 미리 등록해 둘 수 있습니다. 그 사람이 나중에 가입하면 첫 로그인 시 자동으로 멤버로 추가됩니다.")
    add_step(doc, 1, "[이름·이메일로 권한 부여] 영역에서 이름 입력 (조직 회원 자동완성 사용 가능).")
    add_step(doc, 2, "이메일은 선택 (가입 시 매칭 정확도가 높아짐).")
    add_step(doc, 3, "역할 선택 → [추가] — 가입 대기 영역에 미가입 배지로 표시됩니다.")

    add_image_placeholder(
        doc, 21,
        "[이름·이메일로 권한 부여] 영역 + 그 위에 가입 대기 미가입자 1~2명이 노란 카드로 표시된 모습."
    )

    add_heading(doc, "10.4 초대 링크", level=2)
    add_step(doc, 1, "공유 모달에서 [초대 링크 생성].")
    add_step(doc, 2, "링크 복사 후 팀원에게 전달.")
    add_step(doc, 3, "팀원이 링크 접속 + 로그인 → 자동으로 편집 권한으로 추가됩니다(7일 유효).")

    add_tip(doc, "초대 링크는 클릭한 사람을 무조건 편집 권한으로 추가하므로 외부 공유 시 주의하세요. 보기 권한만 주려면 사전 등록 방식을 쓰세요.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 11. 환경설정
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "11. 환경설정", level=1)
    add_para(doc, "헤더 ⋮ → [환경설정]. 탭별로 다른 항목을 관리합니다.")

    add_heading(doc, "11.1 기본 설정", level=2)
    add_bullet(doc, "웹 타이틀 / 크리티컬 패스 표시 / 셀 텍스트 줄바꿈 / WBS ID 접두사")

    add_heading(doc, "11.2 표 컬럼", level=2)
    add_bullet(doc, "기본 컬럼 표시·숨김 토글, 너비 조정")
    add_bullet(doc, "[컬럼 추가] — 사용자 정의 컬럼 (예: 발주처, 산출물 코드)")
    add_bullet(doc, "v0.4.142~ 일반 회원도 사용자 정의 컬럼 추가/수정/삭제 가능. 모든 사용자에게 공유됩니다.")

    add_heading(doc, "11.3 상태/진척도", level=2)
    add_bullet(doc, "상태 정의 추가·이름·색·기본 진척도 설정 (관리자 전용)")
    add_bullet(doc, "상태 ↔ 진척도 연동 토글 — ON이면 상태 바꿀 때 진척도 자동 설정")

    add_heading(doc, "11.4 프로젝트 기간", level=2)
    add_bullet(doc, "프로젝트별 시작일·종료일 한꺼번에 일괄 변경 (본인 소유 또는 관리자만)")

    add_image_placeholder(
        doc, 22,
        "환경설정 모달 — 좌측 탭(기본·표 컬럼·상태·프로젝트 기간) + 우측 활성 탭의 입력 영역."
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 12. 데이터 가져오기·내보내기·백업
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "12. 데이터 가져오기·내보내기·백업", level=1)
    add_heading(doc, "12.1 Excel 가져오기", level=2)
    add_step(doc, 1, "헤더 ⋮ → [Excel 가져오기].")
    add_step(doc, 2, "파일 업로드 → 컬럼 매핑 미리보기 → 검증.")
    add_step(doc, 3, "오류 행은 빨간 표시. 정상만 [임포트].")

    add_heading(doc, "12.2 Excel 내보내기", level=2)
    add_step(doc, 1, "헤더 ⋮ → [내보내기] → Excel 옵션.")
    add_step(doc, 2, "대상(현재/다중/전체) + 컬럼 + 트리 표시 옵션 선택 → [다운로드].")

    add_heading(doc, "12.3 JSON 백업·복원 (관리자 전용)", level=2)
    add_step(doc, 1, "헤더 ⋮ → [백업·복원] → [JSON 백업] — 모든 프로젝트·작업·설정·조직도·감사 로그 포함된 JSON 다운로드.")
    add_step(doc, 2, "복원: [JSON 복원] → 파일 선택 → [전체 교체] / [병합] 모드 → 확인.")
    add_step(doc, 3, "[전체 교체]는 현재 데이터를 모두 삭제하고 백업으로 교체. 사전 자동 백업이 별도로 저장됩니다.")

    add_image_placeholder(
        doc, 23,
        "더보기(⋮) 메뉴 펼친 상태 — [백업·복원], [Excel 가져오기], [내보내기], [환경설정], [공유] 등 항목들."
    )

    add_tip(doc, "백업 JSON은 사용자 정보(이메일·관리자 여부)를 포함하지 않습니다. 다른 환경으로 옮길 때 안전합니다.")

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 13. AI / 주간보고
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "13. AI 보조와 주간보고", level=1)
    add_heading(doc, "13.1 AI 분석 — 작업 자동 추출 (관리자 전용)", level=2)
    add_step(doc, 1, "헤더 ⋮ → [AI 분석].")
    add_step(doc, 2, "기획서·메모 등을 텍스트로 붙여넣기.")
    add_step(doc, 3, "[분석] → 예상 토큰 수·비용 표시 → 확인.")
    add_step(doc, 4, "결과(작업 트리 후보) 미리보기에서 항목별 [수락]/[거절] → 적용.")

    add_image_placeholder(
        doc, 24,
        "AI 분석 모달 — 입력 텍스트 영역 + [분석] 버튼 + 결과 미리보기 표."
    )

    add_heading(doc, "13.2 주간 보고", level=2)
    add_step(doc, 1, "헤더 ⋮ → [주간 보고].")
    add_step(doc, 2, "범위 선택 (전체/내 작업) + 프로젝트 범위 (전체/다중).")
    add_step(doc, 3, "이번 주 완료·진행·이슈가 자동 분류되어 표시.")
    add_step(doc, 4, "이슈 추가 — [+ 이슈] 후 담당자 자동완성으로 입력.")
    add_step(doc, 5, "[복사] / [Excel 다운로드] — 보고서 형식으로 출력.")

    add_image_placeholder(
        doc, 25,
        "주간 보고 모달 — 주차 선택 + 범위 토글 + 자동 분류된 표(금주 한 일·차주 계획·이슈) + [복사]/[Excel] 버튼."
    )

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 14. 단축키
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "14. 키보드 단축키 일람", level=1)
    add_heading(doc, "14.1 공통 (어디서나)", level=2)
    add_kbd_table(doc, [
        ("Ctrl+Z", "Undo (50회)"),
        ("Ctrl+Y / Ctrl+Shift+Z", "Redo"),
        ("Ctrl+K", "검색 모달"),
        ("F11", "풀스크린 토글"),
        ("Esc", "모달·메뉴 닫기"),
    ])

    add_heading(doc, "14.2 표 (WBSTable)", level=2)
    add_kbd_table(doc, [
        ("F2 / Enter (비-편집)", "셀 편집 진입"),
        ("Enter (작업명 편집 중)", "같은 레벨 형제 새 작업"),
        ("Shift+Enter", "현재 행 위에 새 작업"),
        ("Insert", "기준 행의 자식 추가"),
        ("Shift+Insert", "기준 행 위에 형제 추가"),
        ("Tab / Shift+Tab", "들여쓰기 / 내어쓰기 (정렬·필터 OFF 시)"),
        ("Alt+↑ / Alt+↓", "형제 순서 위/아래 이동"),
        ("Shift+→ / Shift+←", "트리 펼치기 / 접기"),
        ("Ctrl+Alt+1~9", "전체 N 레벨까지 펼치기"),
        ("Space", "행 체크박스 토글"),
        ("Delete", "선택 행 삭제"),
        ("Ctrl+C / X / V", "복사 / 잘라내기 / 붙여넣기"),
        ("←/→ (편집 중)", "셀 내 커서 이동 (편집 유지)"),
        ("←/→ (비-편집)", "같은 행에서 컬럼 이동 (행 끝/처음에서 자동 wrap)"),
    ])

    add_heading(doc, "14.3 마인드맵", level=2)
    add_kbd_table(doc, [
        ("Tab", "선택 노드의 자식 추가"),
        ("Ctrl+Enter", "자식 추가"),
        ("Enter", "형제 추가"),
        ("Delete", "선택 노드 삭제"),
    ])

    doc.add_page_break()

    # ────────────────────────────────────────────────────────────────────
    # 15. FAQ
    # ────────────────────────────────────────────────────────────────────
    add_heading(doc, "15. 자주 묻는 질문", level=1)

    add_heading(doc, "Q1. 가입 후 클라우드 동기화가 안 됩니다.", level=2)
    add_para(doc, "관리자 승인 대기 상태입니다. 관리자에게 승인을 요청하거나, 그 동안 로컬 모드로 시범 사용하세요. 승인 후 첫 로그인 시 자동 동기화됩니다.")

    add_heading(doc, "Q2. 부모 작업의 시작/종료일이 자식 변경 시 자동으로 바뀝니다.", level=2)
    add_para(doc, "기본 정책상 부모 일정은 자식 작업의 min/max로 동기화됩니다. 부모 일정을 자식과 무관하게 고정하려면 부모 셀에서 직접 시작일·종료일을 편집하세요. 그 시점에 자동으로 잠금이 추가되어 이후 자식 변경에도 부모 값이 보존됩니다.")

    add_heading(doc, "Q3. 'Ctrl+Z'를 너무 많이 눌렀습니다. 더 되돌리고 싶어요.", level=2)
    add_para(doc, "Undo는 최대 50회까지 보관됩니다. 그 이상은 [백업·복원]에서 사전에 다운로드해 둔 JSON 백업으로 복원해야 합니다.")

    add_heading(doc, "Q4. 다른 사용자가 동시에 같은 셀을 편집하면?", level=2)
    add_para(doc, "동시 편집은 시각적으로만 표시됩니다(다른 사용자 색상의 외곽선). 저장은 마지막이 우선(LWW)이며, 명시적 충돌 해석 UI는 없습니다. 같은 행의 다른 셀은 자유롭게 동시 편집 가능합니다.")

    add_heading(doc, "Q5. 미가입자에게 권한을 미리 줬는데 자동 권한 부여가 안 됩니다.", level=2)
    add_para(doc, "[가입 시 자동 권한 부여]는 그 사람이 회원가입한 직후 첫 로그인 시 동작합니다. 가입 시 사용한 이메일이 사전 등록한 이메일과 일치해야 매칭됩니다. 이메일이 없으면 이름이 정확히 일치해야 합니다(권장: 이메일도 함께 등록).")

    add_heading(doc, "Q6. 회원 체험 모드가 무엇인가요?", level=2)
    add_para(doc, "관리자가 일반 회원 시점으로 화면을 보는 임시 모드입니다. 헤더 ⋮ → [회원 체험] 토글로 켜고, 본인 소유 프로젝트는 여전히 편집 가능하지만 환경설정·조직도 등 관리자 전용 UI는 숨겨집니다. 탭을 닫으면 자동 해제됩니다.")

    doc.add_page_break()

    # ── 끝 페이지 ─────────────────────────────────────────────────────────
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = closing.add_run("\n\n\n— 끝 —\n\n버전 정보·문의는 헤더 우상단 사용자 메뉴를 참고하세요.")
    rc.italic = True
    rc.font.size = Pt(11)
    rc.font.name = "맑은 고딕"

    out_dir = Path("docs/manual")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "사용자매뉴얼.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    p = build()
    print(f"Generated: {p}")
